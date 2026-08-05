"""
End-to-end smoke test for the 6 module-5/6 feature decisions.

Runs in-process against the REAL Postgres (travel_experiment) using
httpx.ASGITransport — no port needed. LLM is forced to MOCK mode so the
agent-chat / behavior-log test is deterministic and needs no network.

Features under test:
  ① reminder_correct  -> researcher manual judgement (set + read back)
  ② quality_score     -> 1-10 researcher score (set + read back)
  ③ plaintext storage -> scores stored unencrypted (DB read asserts bool/int)
  ④ export/all?group= -> group filter + questionnaire score columns
  ⑤ submission detail -> reminders/email + Word(.docx) download
  ⑥ behavior log       -> input_content + ai_response captured on agent chat
"""
import asyncio
import os
import sys
import uuid
from datetime import datetime

sys.path.insert(0, os.path.dirname(__file__))

# Force MOCK LLM (deterministic, no external network from sandbox)
os.environ["LLM_API_KEY"] = ""

from sqlalchemy import select, delete

from app.database import async_session_factory
from app.deps import get_db
from app.models.user import User, ExperimentGroup, UserStatus
from app.models.task_submission import TaskSubmission
from app.models.reminder import Reminder
from app.models.questionnaire import QuestionnaireItem, QuestionnaireResponse, QuestionType
from app.models.admin_score import AdminScore
from app.models.behavior_log import BehaviorLog, ActionType
from app.models.chat_message import ChatMessage
from app.services.auth_service import hash_password
from app.main import app
import httpx

UPLOADS_DIR = os.path.join(os.path.dirname(__file__), "uploads")
TAG = uuid.uuid4().hex[:6]
PW = "Smoke@2026"
results = []


def record(feat, name, ok, detail=""):
    results.append((feat, ok, detail))
    print(f"[{'PASS' if ok else 'FAIL'}] {feat:<10} {name} :: {detail}")


async def get_db_override():
    async with async_session_factory() as s:
        yield s


app.dependency_overrides[get_db] = get_db_override

CREATED = {"users": [], "items": [], "docx_path": None}


async def make_user(email, group):
    async with async_session_factory() as s:
        u = User(
            email=email,
            password_hash=hash_password(PW),
            status=UserStatus.CONSENTED,
            group=group,
        )
        s.add(u)
        await s.commit()
        await s.refresh(u)
        CREATED["users"].append(str(u.id))
        return str(u.id)


async def main():
    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
        # ---------- real ADMIN login ----------
        r = await client.post("/api/auth/admin/login",
                               json={"username": "admin", "password": "admin123"})
        assert r.status_code == 200, f"admin login failed {r.status_code} {r.text}"
        admin_token = r.json()["access_token"]
        ah = {"Authorization": f"Bearer {admin_token}"}
        record("auth", "admin real login", True, "got admin JWT")

        # ---------- create test users ----------
        h_uid = await make_user(f"e2e_h_{TAG}@example.com", ExperimentGroup.H)
        soa_uid = await make_user(f"e2e_soa_{TAG}@example.com", ExperimentGroup.SOA)

        # ---------- seed H user extras (submission + reminder + questionnaire) ----------
        async with async_session_factory() as s:
            sub = TaskSubmission(
                user_id=h_uid, task1_search=True, task2_document=True,
                task3_reminder=True, task4_email=True,
                email_status="sent", email_recipient="dest@example.com",
            )
            s.add(sub)
            rem = Reminder(user_id=h_uid,
                           reminder_datetime=datetime(2026, 8, 1, 9, 0),
                           content="西湖门票预约")
            s.add(rem)
            trust = QuestionnaireItem(construct="trust", question_text="信任度",
                                      question_type=QuestionType.likert,
                                      scale_level=5, sort_order=1, is_active=True)
            s.add(trust)
            await s.commit()
            await s.refresh(trust)
            CREATED["items"].append(str(trust.id))
            s.add(QuestionnaireResponse(user_id=h_uid, item_id=trust.id, response_value="4"))
            mc = QuestionnaireItem(construct="manipulation_check",
                                   question_text="是否使用AI助理",
                                   question_type=QuestionType.choice,
                                   scale_level=0, sort_order=99, is_active=True)
            s.add(mc)
            await s.commit()
            await s.refresh(mc)
            CREATED["items"].append(str(mc.id))
            s.add(QuestionnaireResponse(user_id=h_uid, item_id=mc.id, response_value="是"))
            await s.commit()

        # ===== FEATURE ① + ② : set_score with quality_score + reminder_correct =====
        payload = {f: 8 for f in ["scenic_score", "historic_score", "rarity_score",
                                    "scale_score", "integrity_score", "fame_score",
                                    "season_score"]}
        payload["eco_score"] = 0
        payload["quality_score"] = 7
        payload["reminder_correct"] = True
        payload["notes"] = "冒烟测试"
        r = await client.post(f"/api/admin/scores/{h_uid}", json=payload, headers=ah)
        ok = r.status_code == 200
        record("①②", "set_score quality=7/reminder=True", ok, f"status={r.status_code}")

        # read back via participant detail
        r = await client.get(f"/api/admin/participants/{h_uid}", headers=ah)
        d = r.json()
        sc = (d.get("scores") or [{}])[0]
        ok1 = sc.get("quality_score") == 7
        ok2 = sc.get("reminder_correct") is True
        record("②", "detail quality_score==7", ok1, f"got {sc.get('quality_score')}")
        record("①", "detail reminder_correct==True", ok2, f"got {sc.get('reminder_correct')}")

        # ===== FEATURE ③ : plaintext storage (DB read, no encryption) =====
        async with async_session_factory() as s:
            row = (await s.execute(
                select(AdminScore).where(AdminScore.user_id == h_uid))).scalar_one_or_none()
            ok = isinstance(row.reminder_correct, bool) and row.reminder_correct is True \
                and isinstance(row.quality_score, int) and row.quality_score == 7
            record("③", "scores stored as plaintext bool/int", ok,
                   f"reminder_correct={row.reminder_correct!r} quality_score={row.quality_score!r}")

        # ===== FEATURE ④ : export/all?group=H with questionnaire columns =====
        r = await client.get("/api/admin/export/all", params={"group": "H"}, headers=ah)
        text = r.text
        header = text.splitlines()[0]
        has_q = "Quality_Score" in header
        has_r = "Reminder_Correct" in header
        row = [l for l in text.splitlines()[1:] if f"e2e_h_{TAG}" in l]
        if row:
            fields = row[0].split(",")
            last = fields[-1]
            mc_ok = last.strip() != ""
            # quality & reminder near tail
            tail = fields[-7:]
            record("④", "export header has Quality_Score", has_q, "")
            record("④", "export header has Reminder_Correct", has_r, "")
            record("④", "manipulation_check col non-empty (no tail bug)", mc_ok,
                   f"last_col={last!r}")
            record("④", "H user present in group=H export", True,
                   f"tail={tail}")
        else:
            record("④", "H user present in group=H export", False, "row not found")

        # ===== FEATURE ⑤ : submission detail + Word(.docx) download =====
        # detail already fetched; check reminders + submission email
        rems = d.get("reminders") or []
        sub_d = d.get("submission") or {}
        ok_rem = len(rems) >= 1 and rems[0].get("content") == "西湖门票预约"
        ok_mail = sub_d.get("email_status") == "sent" and sub_d.get("email_recipient") == "dest@example.com"
        record("⑤", "detail shows reminder content", ok_rem, f"reminders={rems}")
        record("⑤", "detail shows sent email", ok_mail, f"email={sub_d.get('email_status')}/{sub_d.get('email_recipient')}")

        # generate a REAL docx and wire it to the submission
        from docx import Document
        os.makedirs(UPLOADS_DIR, exist_ok=True)
        docx_path = os.path.join(UPLOADS_DIR, f"e2e_{TAG}.docx")
        doc = Document()
        doc.add_heading("杭州三日游行程规划（冒烟测试）", 0)
        doc.add_paragraph("第一天：西湖 → 河坊街")
        doc.save(docx_path)
        CREATED["docx_path"] = docx_path
        async with async_session_factory() as s:
            sub2 = (await s.execute(
                select(TaskSubmission).where(TaskSubmission.user_id == h_uid))).scalar_one_or_none()
            sub2.docx_file_path = docx_path
            await s.commit()
        r = await client.get(f"/api/admin/participants/{h_uid}/docx", headers=ah)
        body = r.content
        ok_doc = (r.status_code == 200
                  and "wordprocessingml" in (r.headers.get("content-type") or "")
                  and body[:2] == b"PK")
        record("⑤", "Word .docx download (200 + docx mime + zip)", ok_doc,
               f"status={r.status_code} ctype={r.headers.get('content-type')} bytes={len(body)}")

        # ===== FEATURE ⑥ : behavior log captures input + ai_response =====
        # real participant login for SOA user
        r = await client.post("/api/auth/login",
                              json={"email": f"e2e_soa_{TAG}@example.com", "password": PW})
        assert r.status_code == 200, f"soa login failed {r.status_code} {r.text}"
        user_token = r.json()["access_token"]
        uh = {"Authorization": f"Bearer {user_token}"}
        record("auth", "participant real login (SOA)", True, "got user JWT")

        msg = "你好，我想测试一下对话功能是否正常。"
        async with client.stream("POST", "/api/agent/chat",
                                 json={"agent_id": "soa", "message": msg},
                                 headers=uh) as resp:
            events = []
            async for line in resp.aiter_lines():
                if line.startswith("event:") or line.startswith("data:"):
                    events.append(line)
            last_event = events[-1] if events else ""
            record("⑥", "agent chat SSE streams to 'done'", "event: done" in "\n".join(events),
                   f"events={len(events)} last={last_event[:40]}")

        # read back the behavior log for this user
        async with async_session_factory() as s:
            logs = (await s.execute(
                select(BehaviorLog).where(BehaviorLog.user_id == soa_uid)
                .order_by(BehaviorLog.timestamp.desc()))).scalars().all()
            log = next((l for l in logs if l.action_type == ActionType.AGENT_MESSAGE), None)
            ok_in = log is not None and log.input_content == msg
            ok_ai = log is not None and bool(log.ai_response) and len(log.ai_response.strip()) > 0
            record("⑥", "behavior_log.input_content == user msg", ok_in,
                   f"input={getattr(log,'input_content',None)!r}")
            record("⑥", "behavior_log.ai_response non-empty", ok_ai,
                   f"ai_len={len(log.ai_response) if log else 0} action={log.action_type.value if log else None}")


async def cleanup():
    async with async_session_factory() as s:
        for uid in CREATED["users"]:
            await s.execute(delete(BehaviorLog).where(BehaviorLog.user_id == uid))
            await s.execute(delete(ChatMessage).where(ChatMessage.user_id == uid))
            await s.execute(delete(AdminScore).where(AdminScore.user_id == uid))
            await s.execute(delete(Reminder).where(Reminder.user_id == uid))
            await s.execute(delete(QuestionnaireResponse).where(QuestionnaireResponse.user_id == uid))
            await s.execute(delete(TaskSubmission).where(TaskSubmission.user_id == uid))
            await s.execute(delete(User).where(User.id == uid))
        for iid in CREATED["items"]:
            await s.execute(delete(QuestionnaireItem).where(QuestionnaireItem.id == iid))
        await s.commit()
    if CREATED["docx_path"] and os.path.exists(CREATED["docx_path"]):
        os.remove(CREATED["docx_path"])
    print("CLEANUP done")


if __name__ == "__main__":
    async def run_all():
        try:
            await main()
        finally:
            await cleanup()
            from app.database import engine
            await engine.dispose()

    asyncio.run(run_all())
    fails = [r for r in results if not r[1]]
    print("\n==================== SUMMARY ====================")
    for feat, ok, detail in results:
        print(f"  [{'PASS' if ok else 'FAIL'}] {feat:<10} {detail}")
    print(f"TOTAL {len(results)}  PASS {len(results)-len(fails)}  FAIL {len(fails)}")
    sys.exit(1 if fails else 0)
