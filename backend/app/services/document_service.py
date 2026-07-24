"""Document generation service using python-docx."""
from html.parser import HTMLParser

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class _HtmlToDocx(HTMLParser):
    """Minimal HTML -> docx converter for the rich-text editor output.

    Supports: h1/h2/h3 headings, p/div paragraphs, br line breaks,
    ul/ol/li lists, b/strong / i/em / u inline formatting, and tables.
    """

    def __init__(self, doc):
        super().__init__(convert_charrefs=True)
        self.doc = doc
        self.cur = None  # current paragraph object
        self.bold = False
        self.italic = False
        self.underline = False
        self.list_mode = None  # 'ul' | 'ol'
        # table state
        self.table = None  # {'rows': [...]}
        self.row = None
        self.cell_text = []

    # ---- helpers ----
    def _ensure_para(self):
        if self.cur is None:
            self.cur = self.doc.add_paragraph()

    def _flush_table(self):
        rows = self.table["rows"]
        if not rows:
            return
        ncols = max((len(r) for r in rows), default=0)
        if ncols == 0:
            return
        t = self.doc.add_table(rows=len(rows), cols=ncols)
        try:
            t.style = "Table Grid"
        except Exception:
            pass
        for i, r in enumerate(rows):
            for j in range(ncols):
                t.cell(i, j).text = r[j] if j < len(r) else ""

    # ---- tags ----
    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        if tag in ("h1", "h2", "h3"):
            level = {"h1": 0, "h2": 1, "h3": 2}[tag]
            self.cur = self.doc.add_heading("", level=level)
        elif tag in ("p", "div"):
            self.cur = self.doc.add_paragraph()
        elif tag == "br":
            if self.table is not None and self.row is not None:
                self.cell_text.append("\n")
            elif self.cur is not None and self.cur.text:
                self.cur = self.doc.add_paragraph()
        elif tag in ("b", "strong"):
            self.bold = True
        elif tag in ("i", "em"):
            self.italic = True
        elif tag == "u":
            self.underline = True
        elif tag == "ul":
            self.list_mode = "ul"
        elif tag == "ol":
            self.list_mode = "ol"
        elif tag == "li":
            style = "List Bullet" if self.list_mode == "ul" else "List Number"
            self.cur = self.doc.add_paragraph(style=style)
        elif tag == "table":
            self.table = {"rows": []}
            self.row = None
        elif tag == "tr":
            self.row = []
            self.cell_text = []
        elif tag in ("td", "th"):
            self.cell_text = []
        # unknown tags are ignored but their text still flows through

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in ("b", "strong"):
            self.bold = False
        elif tag in ("i", "em"):
            self.italic = False
        elif tag == "u":
            self.underline = False
        elif tag == "ul":
            self.list_mode = None
        elif tag == "ol":
            self.list_mode = None
        elif tag in ("td", "th"):
            self.row.append("".join(self.cell_text).strip())
            self.cell_text = []
        elif tag == "tr":
            if self.row is not None:
                self.table["rows"].append(self.row)
            self.row = None
        elif tag == "table":
            self._flush_table()
            self.table = None
            self.row = None
        elif tag in ("p", "div", "li", "h1", "h2", "h3"):
            # block-level close: force next text to start a new block
            self.cur = None

    def handle_data(self, data):
        if self.table is not None and self.row is not None:
            self.cell_text.append(data)
            return
        self._ensure_para()
        run = self.cur.add_run(data)
        run.bold = self.bold
        run.italic = self.italic
        run.underline = self.underline


def generate_docx(file_path: str, title: str, content: str, format: str = "text"):
    """Generate a formatted .docx file.

    format='text': parse plain text with lightweight markdown-ish markers
                   (# heading, - bullet, N. numbered list).
    format='html': parse HTML produced by the rich-text editor.
    """
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)

    # Title
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if format == "html":
        parser = _HtmlToDocx(doc)
        parser.feed(content or "")
    else:
        # Parse and add content (legacy plain-text behaviour)
        for line in (content or "").split("\n"):
            line = line.strip()
            if not line:
                continue

            if line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith(tuple("0123456789")) and ". " in line:
                idx = line.index(". ")
                doc.add_paragraph(line[idx + 2:], style="List Number")
            else:
                doc.add_paragraph(line)

    doc.save(file_path)
